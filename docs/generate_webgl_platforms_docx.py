#!/usr/bin/env python3
"""Generate WebGL Distribution Platforms checklist as .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title ──
t = doc.add_heading('WebGL Game Distribution — All 41 Platforms', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs:
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('IntelliVerseX SDK-Powered Game  •  March 2026  •  Release Tracker & Checklist')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# ── Helper ──
def heading(text, level, color):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = color

def table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri + 1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

def checklist(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f'☐  {item}')
        r.font.size = Pt(10)

GREEN  = RGBColor(0x1B, 0x7A, 0x2B)
BLUE   = RGBColor(0x0D, 0x47, 0xA1)
ORANGE = RGBColor(0xE6, 0x51, 0x00)
PURPLE = RGBColor(0x6A, 0x1B, 0x9A)
DARK   = RGBColor(0x33, 0x33, 0x33)

# ════════════════════════════════════════
# SUMMARY
# ════════════════════════════════════════
heading('Platform Summary', 1, DARK)
p = doc.add_paragraph()
r = p.add_run('41 platforms across 8 tiers where your WebGL/HTML5 game can be released.')
r.font.size = Pt(11)

table(
    ['Tier', 'Platforms', 'Count', 'Best For'],
    [
        ['Tier 1', 'Major Game Portals', '7', 'High traffic, established audience'],
        ['Tier 2', 'Distribution Networks', '4', 'One upload → thousands of publishers'],
        ['Tier 3', 'Self-Publishing / Indie', '5', 'Full control, community building'],
        ['Tier 4', 'Social / Messaging', '7', 'Viral growth, social features'],
        ['Tier 5', 'App Stores (PWA/TWA)', '4', 'Mobile reach without native build'],
        ['Tier 6', 'Gaming Platforms', '7', 'Niche communities, casual players'],
        ['Tier 7', 'Web3 / Crypto', '3', 'Token rewards, NFT integration'],
        ['Tier 8', 'Embedding / White-Label', '4', 'Direct revenue, licensing deals'],
    ]
)

# ════════════════════════════════════════
# TIER 1
# ════════════════════════════════════════
heading('Tier 1 — Major Game Portals', 1, GREEN)

table(
    ['#', 'Platform', 'Monthly Users', 'Revenue', 'SDK', 'Priority'],
    [
        ['1', 'Poki', '80M+', 'Ad rev share', 'Poki SDK', 'P0'],
        ['2', 'CrazyGames', '40M+', '70% ad rev share', 'CrazyGames SDK', 'P0'],
        ['3', 'Y8', '30M+', 'Ad rev share', 'Y8 SDK', 'P2'],
        ['4', 'Newgrounds', '10M+', 'Ads + Supporter', 'None', 'P1'],
        ['5', 'Kongregate', '8M+', 'Ads + Virtual Goods', 'Kongregate SDK', 'P2'],
        ['6', 'Addicting Games', '5M+', 'Ad rev share', 'None', 'P2'],
        ['7', 'Armor Games', '3M+', 'Ad rev share', 'None', 'P2'],
    ]
)

heading('Poki — Checklist', 2, DARK)
checklist([
    'Register developer account at developers.poki.com',
    'Integrate Poki SDK (ad breaks, gameplay tracking)',
    'Submit game for review (curated — quality bar)',
    'Prepare 16:9 screenshots and thumbnail',
    'Test on mobile + desktop browsers',
    'Optimize initial load < 5 MB',
])

heading('CrazyGames — Checklist', 2, DARK)
checklist([
    'Register at developer.crazygames.com',
    'Integrate CrazyGames SDK (ads, analytics)',
    'Upload WebGL build as ZIP',
    'Prepare cover image (1200×675) and screenshots',
    'Test on Chrome, Firefox, Safari, Edge',
    'Submit for review',
])

heading('Newgrounds — Checklist', 2, DARK)
checklist([
    'Create Newgrounds developer account',
    'Upload HTML5 game (no SDK required)',
    'Add game description, tags, and thumbnail',
    'Engage with community (comments, forums)',
])

heading('Y8 / Kongregate / Addicting Games / Armor Games', 2, DARK)
checklist([
    'Register developer accounts on each portal',
    'Upload WebGL build (ZIP or iframe URL)',
    'Integrate portal-specific SDK if required',
    'Prepare assets: thumbnail, screenshots, description',
    'Submit and await review',
])

# ════════════════════════════════════════
# TIER 2
# ════════════════════════════════════════
heading('Tier 2 — Distribution Networks', 1, GREEN)

table(
    ['#', 'Platform', 'Reach', 'Revenue', 'SDK', 'Priority'],
    [
        ['8', 'GameDistribution', '7,500+ publishers', '50% rev share', 'GD SDK', 'P1'],
        ['9', 'GameMonetize', '3,500+ publishers', '45% (90% self-pub)', 'GM SDK', 'P2'],
        ['10', 'GamePix', '1,000+ publishers', 'Rev share', 'GamePix SDK', 'P2'],
        ['11', 'Softgames', '500M+ reach', 'Rev share', 'Softgames SDK', 'P2'],
    ]
)

heading('GameDistribution — Checklist', 2, DARK)
checklist([
    'Register at gamedistribution.com/developers',
    'Integrate GameDistribution SDK (ads, events)',
    'Upload game build',
    'Game auto-distributed to 7,500+ publisher sites',
    'Track analytics via GD dashboard',
])

heading('GameMonetize — Checklist', 2, DARK)
checklist([
    'Register at gamemonetize.com/developers',
    'Integrate GameMonetize SDK (5 min integration)',
    'Upload game (supports HTML5, Construct, Unity WebGL)',
    'Optionally self-publish on own site for 90% total share',
    'Payout via PayPal or USDT (Net 30, $30 min)',
])

# ════════════════════════════════════════
# TIER 3
# ════════════════════════════════════════
heading('Tier 3 — Self-Publishing & Indie', 1, BLUE)

table(
    ['#', 'Platform', 'Model', 'SDK', 'Priority'],
    [
        ['12', 'itch.io', 'Free / PWYW / Paid', 'None', 'P0'],
        ['13', 'Game Jolt', 'Free + ads', 'None', 'P1'],
        ['14', 'Miniplay', 'Ad rev share', 'None', 'P3'],
        ['15', 'SilverGames', 'Ad rev share', 'None', 'P3'],
        ['16', 'WebGamer', 'Ad-free / iframe', 'None', 'P3'],
    ]
)

heading('itch.io — Checklist', 2, DARK)
checklist([
    'Create itch.io developer account',
    'Upload WebGL build as ZIP (index.html at root)',
    'Configure embed settings (viewport size, fullscreen)',
    'Set pricing (free / PWYW / paid)',
    'Write game page (description, screenshots, GIFs)',
    'Tag appropriately for discovery (quiz, trivia, webgl)',
    'Join relevant game jams for visibility',
])

heading('Game Jolt — Checklist', 2, DARK)
checklist([
    'Create Game Jolt developer account',
    'Upload HTML5 build',
    'Write game description and add screenshots',
    'Engage with community features (followers, comments)',
])

# ════════════════════════════════════════
# TIER 4
# ════════════════════════════════════════
heading('Tier 4 — Social & Messaging Platforms', 1, BLUE)

table(
    ['#', 'Platform', 'Users', 'Revenue', 'SDK', 'Priority'],
    [
        ['17', 'Telegram Mini Apps', '900M+', 'IAP + Ads', 'Bot API + Mini Apps SDK', 'P0'],
        ['18', 'Discord Activities', '200M+', 'IAP + Subs', 'Embedded App SDK', 'P1'],
        ['19', 'Facebook Instant Games', '2B+', 'Ads + IAP', 'FB Instant Games SDK', 'P1'],
        ['20', 'LINE Games', '200M+ (Asia)', 'IAP', 'LINE SDK', 'P3'],
        ['21', 'WeChat Mini Games', '1.3B+ (China)', 'IAP', 'WeChat SDK', 'P3'],
        ['22', 'Snapchat Games', '400M+', 'Ads', 'Snap SDK', 'P3'],
        ['23', 'VK Games', '100M+ (CIS)', 'Ads + IAP', 'VK Bridge SDK', 'P3'],
    ]
)

heading('Telegram Mini Apps — Checklist', 2, DARK)
checklist([
    'Create Telegram Bot via @BotFather',
    'Set up Mini App (Web App) URL pointing to your WebGL build',
    'Integrate Telegram Mini Apps SDK (auth, payments, sharing)',
    'Add high score / leaderboard via Bot Gaming API',
    'Enable full-screen mode for immersive gameplay',
    'Test viral sharing (forward game to chats / groups)',
    'Submit to Telegram Mini App catalog',
    'Consider TON token integration for rewards (Tier 7)',
])

heading('Discord Activities — Checklist', 2, DARK)
checklist([
    'Register Discord application at discord.com/developers',
    'Integrate Discord Embedded App SDK',
    'Configure activity for voice channel play',
    'Set up IAP / subscriptions if applicable',
    'Test multiplayer in voice channels',
    'Submit for review',
])

heading('Facebook Instant Games — Checklist', 2, DARK)
checklist([
    'Create Facebook App at developers.facebook.com',
    'Integrate FB Instant Games SDK (ads, social, payments)',
    'Upload WebGL build via App Dashboard',
    'Configure Messenger bot for game discovery',
    'Test social features (challenges, leaderboards, sharing)',
    'Submit for review',
])

# ════════════════════════════════════════
# TIER 5
# ════════════════════════════════════════
heading('Tier 5 — App Stores (PWA / TWA Wrapping)', 1, ORANGE)

table(
    ['#', 'Platform', 'Reach', 'Method', 'Priority'],
    [
        ['24', 'Google Play Store', '2B+ Android', 'Trusted Web Activity (TWA)', 'P2'],
        ['25', 'Microsoft Store', '1B+ Windows', 'Progressive Web App (PWA)', 'P2'],
        ['26', 'Samsung Galaxy Store', 'Samsung devices', 'PWA / TWA', 'P3'],
        ['27', 'Amazon Appstore', 'Fire devices', 'Hosted Web App', 'P3'],
    ]
)

heading('Google Play Store (TWA) — Checklist', 2, DARK)
checklist([
    'Set up TWA using Bubblewrap or PWABuilder',
    'Generate signed APK/AAB from your WebGL URL',
    'Create Google Play Developer account ($25 one-time)',
    'Prepare store listing (screenshots, description, icon)',
    'Upload AAB and submit for review',
    'Ensure Digital Asset Links for TWA verification',
])

heading('Microsoft Store (PWA) — Checklist', 2, DARK)
checklist([
    'Add PWA manifest (manifest.json) to your WebGL game',
    'Register Microsoft Partner Center account',
    'Submit PWA URL or upload MSIX package',
    'Prepare store listing assets',
])

# ════════════════════════════════════════
# TIER 6
# ════════════════════════════════════════
heading('Tier 6 — Gaming-Specific Platforms', 1, ORANGE)

table(
    ['#', 'Platform', 'Type', 'Priority'],
    [
        ['28', 'Now.gg', 'Cloud gaming', 'P2'],
        ['29', 'iogames.space', '.io game portal', 'P3'],
        ['30', 'Lagged.com', 'Casual portal', 'P3'],
        ['31', 'Kizi', 'Family-friendly portal', 'P3'],
        ['32', 'Friv', 'Massive casual portal', 'P3'],
        ['33', 'PlayOnRay', '80% rev share', 'P2'],
        ['34', 'Zihit.com', 'Dev tools + analytics', 'P3'],
    ]
)

heading('Checklist (All Tier 6)', 2, DARK)
checklist([
    'Register developer accounts on target portals',
    'Upload WebGL build (ZIP or iframe URL)',
    'Prepare thumbnail, screenshots, description',
    'Optimize for mobile browsers (touch input, responsive)',
    'Submit for review where applicable',
])

# ════════════════════════════════════════
# TIER 7
# ════════════════════════════════════════
heading('Tier 7 — Web3 / Crypto Platforms', 1, PURPLE)

table(
    ['#', 'Platform', 'Model', 'Priority'],
    [
        ['35', 'TON (Telegram Open Network)', 'Token rewards', 'P2'],
        ['36', 'Gala Games', 'Token + NFT', 'P3'],
        ['37', 'Immutable X', 'NFT assets', 'P3'],
    ]
)

heading('TON / Telegram Crypto — Checklist', 2, DARK)
checklist([
    'Integrate TON Connect SDK for wallet linking',
    'Design token reward mechanics for quiz completion',
    'Set up smart contract for reward distribution',
    'Test on Telegram with TON testnet',
    'Launch on TON mainnet',
])

# ════════════════════════════════════════
# TIER 8
# ════════════════════════════════════════
heading('Tier 8 — Embedding & White-Label', 1, PURPLE)

table(
    ['#', 'Platform', 'Model', 'Priority'],
    [
        ['38', 'Own Website', 'Direct / Ads', 'P0'],
        ['39', 'News / Media Sites', 'License fee', 'P2'],
        ['40', 'Education Portals', 'Free / Licensed', 'P2'],
        ['41', 'Corporate Training', 'License / SaaS', 'P3'],
    ]
)

heading('Own Website — Checklist', 2, DARK)
checklist([
    'Deploy WebGL build to your domain (e.g. play.intelliversex.com)',
    'Add AdSense / ad network for monetization',
    'Set up analytics (GA4, Mixpanel)',
    'SEO optimize game page',
    'Add social sharing buttons',
    'Implement HTTPS and proper CORS headers',
])

heading('Education / Corporate — Checklist', 2, DARK)
checklist([
    'Prepare white-label version (removable branding)',
    'Create licensing terms and pricing',
    'Reach out to education portals (Coolmath, ABCya, PBS Kids)',
    'Prepare corporate demo for L&D departments',
])

# ════════════════════════════════════════
# PRIORITY SUMMARY
# ════════════════════════════════════════
doc.add_page_break()
heading('Release Priority Order — Quiz Game', 1, DARK)
p = doc.add_paragraph()
r = p.add_run('Recommended order based on effort-to-impact ratio for a quiz game:')
r.font.italic = True

table(
    ['Priority', '#', 'Platform', 'Why'],
    [
        ['P0', '38', 'Own Website', 'Full control, best margins, immediate'],
        ['P0', '12', 'itch.io', 'Instant publish, build community, free'],
        ['P0', '2', 'CrazyGames', 'High traffic, great SDK, 70% rev share'],
        ['P0', '1', 'Poki', 'Largest portal, curated quality'],
        ['P0', '17', 'Telegram Mini Apps', 'Viral sharing, 900M users, quiz games thrive'],
        ['P1', '8', 'GameDistribution', 'One upload → 7,500+ publishers'],
        ['P1', '18', 'Discord Activities', 'Social play in voice channels'],
        ['P1', '19', 'Facebook Instant Games', 'Social viral loop, quiz games do well'],
        ['P1', '4', 'Newgrounds', 'Indie-friendly, no SDK overhead'],
        ['P1', '13', 'Game Jolt', 'Community-driven, game jams'],
        ['P2', '24', 'Google Play (TWA)', 'Wrap as Android app, mobile reach'],
        ['P2', '25', 'Microsoft Store (PWA)', 'Windows store listing'],
        ['P2', '9', 'GameMonetize', 'Additional publisher network'],
        ['P2', '3', 'Y8', 'Established audience'],
        ['P2', '35', 'TON Crypto', 'Token rewards for quiz participation'],
        ['P2', '33', 'PlayOnRay', '80% rev share'],
        ['P2', '39', 'News / Media Sites', 'Licensing revenue'],
        ['P2', '40', 'Education Portals', 'Quiz games fit perfectly'],
        ['P3', '5-7', 'Kongregate / Addicting / Armor', 'Secondary portals'],
        ['P3', '10-11', 'GamePix / Softgames', 'Additional networks'],
        ['P3', '20-23', 'LINE / WeChat / Snap / VK', 'Regional expansion'],
        ['P3', '26-27', 'Samsung / Amazon Store', 'Device-specific stores'],
        ['P3', '14-16', 'Miniplay / Silver / WebGamer', 'Niche portals'],
        ['P3', '36-37', 'Gala / Immutable X', 'Web3 gaming'],
        ['P3', '41', 'Corporate Training', 'Enterprise licensing'],
    ]
)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Generated March 2026 — IntelliVerseX SDK Engineering Team  •  41 Platforms Total')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.font.italic = True

out = os.path.join(os.path.dirname(__file__), 'IntelliVerseX_WebGL_Distribution_Platforms.docx')
doc.save(out)
print(f'Saved to {out}')
