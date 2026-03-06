#!/usr/bin/env python3
"""Generate the SDK Release Checklist as a .docx file."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title ──
title = doc.add_heading('IntelliVerseX SDK — Release Checklist', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 5.1.0  •  March 2026  •  8 Platform SDKs  •  17 Distribution Channels')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ── Helper functions ──
def add_phase_heading(text, color_rgb):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = color_rgb

def add_channel_heading(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_checklist(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'☐  {item}')
        run.font.size = Pt(11)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            row_cells[ci].text = val
            for p in row_cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

# ── Platform Readiness Summary ──
doc.add_heading('Platform Readiness Summary', level=1)
add_table(
    ['Platform', 'Language', 'Readiness', 'Tests', 'Examples', 'Status'],
    [
        ['Unity / .NET', 'C#', '93%', 'EditMode + PlayMode', '9 categories', 'Production Ready'],
        ['JavaScript / TS', 'TypeScript', '92%', '12 vitest tests', 'Browser + Node.js', 'Production Ready'],
        ['Unreal Engine', 'C++ / Blueprints', '91%', 'Example GameMode', 'ExampleGameMode', 'Production Ready'],
        ['Godot Engine', 'GDScript', '91%', '25 GUT tests', 'basic_example.gd', 'Production Ready'],
        ['C / C++ Native', 'C++', '91%', '14 assert tests', 'main.cpp', 'Production Ready'],
        ['Java / Android', 'Java', '91%', '20 JUnit 5 tests', 'BasicExample.java', 'Production Ready'],
        ['Defold', 'Lua', '90%', '15 tests', 'basic_example.lua', 'Production Ready'],
        ['Cocos2d-x', 'C++', '90%', 'Manual', 'ExampleScene', 'Production Ready'],
        ['Flutter / Dart', 'Dart', '0%', 'N/A', 'N/A', 'Not Started'],
        ['Web3', 'TypeScript', '0%', 'N/A', 'N/A', 'Not Started'],
    ]
)

# ════════════════════════════════════════════════════════════════
# PHASE 1
# ════════════════════════════════════════════════════════════════
add_phase_heading('Phase 1 — Ship Now (Week of Mar 10)', RGBColor(0x1B, 0x7A, 0x2B))

add_channel_heading('1. GitHub Releases — All 8 SDKs')
add_checklist([
    'Clean up repo — remove .bak files and build artifacts',
    'Add license headers to all source files',
    'Write CHANGELOG.md',
    'Create release tags (v5.1.0) per SDK',
    'Attach platform-specific zip bundles to release',
])

add_channel_heading('2. Unity Asset Store — Unity / .NET (93%)')
add_checklist([
    'Set up Asset Store Publisher account',
    'Prepare 5+ screenshots and key art',
    'Record demo video (2–3 min)',
    'Write store listing copy (description, features, requirements)',
    'Submit for review',
])

add_channel_heading('3. npm Registry — JavaScript / TypeScript (92%)')
add_checklist([
    'Set up npm account / org (@intelliversex)',
    'Run npm publish --dry-run to validate package',
    'Optimize bundle size',
    'Publish @intelliversex/sdk',
])

add_channel_heading('4. Nakama Community — All Platforms')
add_checklist([
    'Write showcase post for Nakama forum',
    'Prepare sample Nakama server configs',
    'Submit as community plugin',
])

add_channel_heading('5. Reddit / Indie Communities — All Platforms')
add_checklist([
    'Write launch posts for r/gamedev, r/unity3d, r/godot, r/indiegames',
    'Prepare screenshots and GIFs',
    'Schedule posts for optimal timing',
])

# ════════════════════════════════════════════════════════════════
# PHASE 2
# ════════════════════════════════════════════════════════════════
add_phase_heading('Phase 2 — Fast Follow (Weeks 3–4)', RGBColor(0x0D, 0x47, 0xA1))

add_channel_heading('6. OpenUPM — Unity / .NET (93%)')
add_checklist([
    'Register package on openupm.com',
    'Verify openupm add com.intelliversex.sdk install flow',
    'Add install badge to README',
])

add_channel_heading('7. Unreal Marketplace — Unreal Engine (91%)')
add_checklist([
    'Set up Epic Marketplace seller account',
    'Build-test plugin against real UE 5.3 / 5.4',
    'Package plugin per marketplace guidelines',
    'Prepare marketplace listing (icon, screenshots, description)',
    'Submit for review',
])

add_channel_heading('8. Godot Asset Library — Godot Engine (91%)')
add_checklist([
    'Build-test addon against Godot 4.2+',
    'Submit to Godot Asset Library',
    'Verify install via AssetLib browser',
])

add_channel_heading('9. Maven Central — Java / Android (91%)')
add_checklist([
    'Set up Sonatype OSSRH account',
    'Generate GPG signing key',
    'Configure Gradle maven-publish plugin',
    'Publish com.intelliversex:sdk',
])

add_channel_heading('10. Itch.io Tools — All Platforms')
add_checklist([
    'Create itch.io developer page',
    'Package per-platform SDK bundles (zip)',
    'Write itch.io page (screenshots, description, docs links)',
])

add_channel_heading('11. Developer Portal — All Platforms')
add_checklist([
    'Build developer.intelliversex.com',
    'SDK downloads page per platform',
    'API documentation (auto-generated)',
    'Getting started guides',
    'Dashboard for API keys / project config',
])

add_channel_heading('12. Product Hunt — Developer Portal')
add_checklist([
    'Prepare Product Hunt launch page',
    'Screenshots + demo video',
    'Schedule launch day (coordinate with portal launch)',
])

# ════════════════════════════════════════════════════════════════
# PHASE 3
# ════════════════════════════════════════════════════════════════
add_phase_heading('Phase 3 — Expand (Weeks 5–8)', RGBColor(0xE6, 0x51, 0x00))

add_channel_heading('13. vcpkg / Conan — C / C++ (91%)')
add_checklist([
    'Write vcpkg portfile',
    'Write Conan recipe',
    'Submit PRs to vcpkg-registry / conan-center-index',
    'Verify vcpkg install intelliversex flow',
])

add_channel_heading('14. CodeCanyon — Unity + JS (92%+)')
add_checklist([
    'Create premium listing',
    'Prepare documentation package',
    'Record demo video',
    'Submit for review',
])

add_channel_heading('15. GameDev Market — Unity + Unreal (91%+)')
add_checklist([
    'Create seller account',
    'Prepare listing (screenshots, description)',
    'Submit packages',
])

# ════════════════════════════════════════════════════════════════
# PHASE 4
# ════════════════════════════════════════════════════════════════
add_phase_heading('Phase 4 — Future', RGBColor(0x6A, 0x1B, 0x9A))

add_channel_heading('16. Pub.dev — Flutter / Dart (0% — needs new SDK)')
add_checklist([
    'Build Flutter SDK from scratch',
    'Integrate with nakama-dart client',
    'Write tests and examples',
    'Publish to pub.dev',
])

add_channel_heading('17. Web3 Platforms — JavaScript (0% — needs backend)')
add_checklist([
    'Build token/NFT reward integration',
    'Integrate Thirdweb or Moralis connectors',
    'Build Web3 wallet linking flow',
])

# ════════════════════════════════════════════════════════════════
# BUILD VERIFICATION TABLE
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
add_phase_heading('Platform Build Verification (Pre-Release)', RGBColor(0x33, 0x33, 0x33))
p = doc.add_paragraph('Complete these build tests before submitting to any marketplace:')
p.runs[0].font.italic = True

add_table(
    ['☐', 'Platform', 'Build Target', 'OS Targets'],
    [
        ['☐', 'Unity', 'Unity 2023 LTS + Unity 6', 'Windows, Mac, Android, iOS, WebGL'],
        ['☐', 'Unreal', 'UE 5.3 / 5.4', 'Win64, Mac, Linux, Android, iOS'],
        ['☐', 'Godot', 'Godot 4.2+', 'Windows, Mac, Linux, Android, iOS, Web'],
        ['☐', 'Defold', 'Defold stable', 'Windows, Mac, Linux, Android, iOS, Web'],
        ['☐', 'Cocos2d-x', 'Cocos2d-x 4.0', 'Windows, Mac, Android, iOS'],
        ['☐', 'JavaScript', 'Node 18/20/22 + Browsers', 'Chrome, Firefox, Safari, Edge'],
        ['☐', 'C/C++', 'CMake 3.14+', 'Ubuntu, macOS, Windows (MSVC/GCC/Clang)'],
        ['☐', 'Java', 'JDK 11/17/21', 'Desktop + Android'],
    ]
)

# ── Channel Summary Table ──
doc.add_heading('Distribution Channel Summary', level=1)
add_table(
    ['#', 'Channel', 'SDK', 'SDK Ready', 'Channel Ready', 'Priority', 'Target'],
    [
        ['1', 'GitHub Releases', 'All 8 SDKs', '90%+', '90%', 'P0', 'Mar 10'],
        ['2', 'Unity Asset Store', 'Unity', '93%', '70%', 'P0', 'Mar 15'],
        ['3', 'npm Registry', 'JavaScript', '92%', '85%', 'P1', 'Mar 18'],
        ['4', 'Nakama Community', 'All', '90%+', '50%', 'P1', 'Mar 12'],
        ['5', 'Reddit / Indie', 'All', '90%+', '30%', 'P1', 'Mar 15'],
        ['6', 'OpenUPM', 'Unity', '93%', '60%', 'P1', 'Mar 20'],
        ['7', 'Developer Portal', 'All', '90%+', '10%', 'P1', 'Apr 1'],
        ['8', 'Unreal Marketplace', 'Unreal', '91%', '50%', 'P2', 'Apr 10'],
        ['9', 'Godot Asset Library', 'Godot', '91%', '55%', 'P2', 'Apr 5'],
        ['10', 'Itch.io Tools', 'All', '90%+', '30%', 'P2', 'Apr 1'],
        ['11', 'Maven Central', 'Java', '91%', '40%', 'P2', 'Apr 15'],
        ['12', 'Product Hunt', 'Portal', '90%+', '5%', 'P2', 'Apr 15'],
        ['13', 'vcpkg / Conan', 'C/C++', '91%', '20%', 'P3', 'May 1'],
        ['14', 'CodeCanyon', 'Unity + JS', '92%+', '20%', 'P3', 'May 1'],
        ['15', 'GameDev Market', 'Unity + UE', '91%+', '10%', 'P3', 'May 15'],
        ['16', 'Pub.dev', 'Flutter', '0%', '0%', 'P3', 'May 15'],
        ['17', 'Web3 Platforms', 'JavaScript', '0%', '0%', 'P4', 'Jun 1'],
    ]
)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Generated March 2026 — IntelliVerseX SDK Engineering Team')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

# ── Save ──
out_path = os.path.join(os.path.dirname(__file__), 'IntelliVerseX_SDK_Release_Checklist.docx')
doc.save(out_path)
print(f'Saved to {out_path}')
